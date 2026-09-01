from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\projek UMKM\gis-umkm-mapping-main")
REFERENCE = Path(r"C:\Users\TUF GAMING\OneDrive\Dokumen\Skripsi Bangun BISINDO.docx")
OUTPUT = ROOT / "Skripsi_Rancang_Bangun_Web_GIS_Zonasi_UMKM_Sulawesi_Utara.docx"
TITLE = (
    "RANCANG BANGUN SISTEM INFORMASI GEOGRAFIS BERBASIS WEB UNTUK "
    "PEMETAAN DAN ANALISIS ZONASI UMKM DI SULAWESI UTARA "
    "MENGGUNAKAN ALGORITMA K-MEANS"
)
NAME = "[NAMA MAHASISWA]"
NIM = "[NIM]"


def set_font(run, size: float = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    start = OxmlElement("w:fldChar")
    start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([start, instruction, end])


def set_cell_style(cell, fill: str | None = None) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", 90), ("start", 100), ("bottom", 90), ("end", 100)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    properties.append(margins)
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "808080")
        borders.append(border)
    properties.append(borders)


def mark_header_row(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


document = Document(REFERENCE)
clear_document_body(document)
section = document.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(4)
section.right_margin = Cm(3)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)
section.different_first_page_header_footer = True

normal = document.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.first_line_indent = Cm(1.25)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for style_name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
    style = document.styles[style_name]
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.line_spacing = 1.5

header = section.header.paragraphs[0]
header.text = ""
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Halaman "), 9)
add_page_number(footer)
first_footer = section.first_page_footer.paragraphs[0]
first_footer.clear()


def add_paragraph(
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent: bool = True,
    bold: bool = False,
    italic: bool = False,
    size: float = 12,
    before: float = 0,
    after: float = 0,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    set_font(paragraph.add_run(text), size, bold, italic)
    return paragraph


def add_centered(text: str, size: float = 12, bold: bool = False, after: float = 0):
    return add_paragraph(
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        size=size,
        bold=bold,
        after=after,
    )


def add_heading(text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    set_font(paragraph.add_run(text), 14 if level == 1 else 12, bold=True)


def add_table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        set_cell_style(cell, "D9EAD3")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].add_run(text), 9, bold=True)
        if widths:
            cell.width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = ""
            set_cell_style(cells[index])
            set_font(cells[index].paragraphs[0].add_run(text), 8.5)
            if widths:
                cells[index].width = Cm(widths[index])


def next_page() -> None:
    document.add_page_break()


def academic_paragraphs(subject: str, evidence: str, implementation: str, interpretation: str) -> list[str]:
    return [
        f"{subject}. {evidence} Pembahasan ditempatkan pada aplikasi Zonasi UMKM Sulawesi Utara yang menghubungkan identitas usaha, kategori, alamat, koordinat, status ketelitian lokasi, dan hasil pengelompokan. Data tabular menjadi sumber administrasi, sedangkan peta memberi konteks ruang untuk membaca kedekatan serta konsentrasi. Keduanya diproses dalam satu alur agar informasi tidak terlepas dari asal dan kualitas datanya.",
        f"Implementasi dilakukan melalui {implementation}. Data dinormalisasi dan diperiksa kelengkapan pasangan lintang-bujurnya. Koordinat analitis digunakan untuk perhitungan, sedangkan koordinat tampilan digunakan untuk marker. Status lokasi membedakan titik terverifikasi, titik perkiraan wilayah, dan data yang belum dapat dipetakan. Pemisahan ini mencegah aplikasi menyampaikan presisi yang tidak dimiliki sumber data.",
        f"Dari sisi pengguna, {interpretation} Halaman publik menyediakan pencarian, filter, popup, informasi klaster, dan UMKM terdekat. Panel pengelola menyediakan autentikasi, pembatasan peran, verifikasi koordinat, impor-ekspor, analisis, snapshot, dan audit. Klaster tetap dipahami sebagai ringkasan kedekatan berdasarkan nilai K dan snapshot data, bukan batas administrasi atau peringkat usaha. Hasilnya digunakan untuk eksplorasi dan prioritas verifikasi yang kemudian dikonfirmasi dengan data lapangan.",
    ]


def add_topic_page(chapter: str, section_title: str, paragraphs: list[str]) -> None:
    add_heading(chapter)
    add_heading(section_title, 2)
    for text in paragraphs:
        add_paragraph(text)
    next_page()


add_centered("SKRIPSI", 14, True, 22)
add_centered(TITLE, 14, True, 55)
add_centered("Disusun:", 12, False, 8)
add_centered(NAME, 12, True)
add_centered(f"NIM. {NIM}", 12, True, 60)
add_centered("POLITEKNIK NEGERI MANADO", 12, True)
add_centered("JURUSAN TEKNIK ELEKTRO", 12, True)
add_centered("PROGRAM STUDI D-IV TEKNIK INFORMATIKA", 12, True)
add_centered("2026", 12, True)
next_page()

add_centered(TITLE, 14, True, 24)
add_centered("SKRIPSI", 14, True, 18)
add_centered("Disusun untuk melengkapi salah satu syarat kelulusan")
add_centered("Program Sarjana Terapan Program Studi D-IV Teknik Informatika")
add_centered("Jurusan Teknik Elektro di Politeknik Negeri Manado", after=30)
add_centered("Disusun:")
add_centered(NAME, 12, True)
add_centered(f"NIM. {NIM}", 12, True, 55)
add_centered("POLITEKNIK NEGERI MANADO", 12, True)
add_centered("JURUSAN TEKNIK ELEKTRO", 12, True)
add_centered("PROGRAM STUDI D-IV TEKNIK INFORMATIKA", 12, True)
add_centered("2026", 12, True)
next_page()

add_centered("HALAMAN PENGESAHAN", 14, True, 24)
add_centered(TITLE, 12, True, 22)
add_centered("Disusun:")
add_centered(NAME, 12, True)
add_centered(f"NIM. {NIM}", 12, True, 24)
add_paragraph(
    "Telah dipertahankan dalam Seminar dan Ujian Skripsi di depan Tim Penguji pada tanggal "
    "................................ dan dinyatakan telah memenuhi syarat sebagai Sarjana Terapan.",
    indent=False,
)
add_centered("Disahkan oleh:", 12, True, 16)
add_table(
    ["Jabatan", "Nama dan tanda tangan"],
    [
        ["Ketua Pelaksana Ujian Skripsi", "................................................"],
        ["Pembimbing I", "................................................"],
        ["Pembimbing II", "................................................"],
    ],
    [6, 8],
)
next_page()

add_centered("PERNYATAAN KEASLIAN TULISAN", 14, True, 22)
add_paragraph("Yang bertanda tangan di bawah ini:", indent=False)
add_table(
    ["Keterangan", "Isi"],
    [["Nama", NAME], ["NIM", NIM], ["Program Studi", "D-IV Teknik Informatika"], ["Jurusan", "Teknik Elektro"]],
    [4, 10],
)
add_paragraph(
    "menyatakan bahwa skripsi ini merupakan hasil karya saya sendiri. Seluruh kutipan, data, dan "
    "rujukan dicantumkan sesuai kaidah ilmiah. Apabila terbukti terdapat pelanggaran akademik, "
    "saya bersedia menerima sanksi sesuai ketentuan yang berlaku."
)
add_centered("Manado, .......................... 2026", after=10)
add_centered("Yang membuat pernyataan,", after=55)
add_centered(NAME, 12, True)
add_centered(f"NIM. {NIM}", 12, True)
next_page()

add_centered("ABSTRAK", 14, True, 20)
add_paragraph(f"{NAME}, {TITLE.lower().capitalize()} (dibimbing oleh Pembimbing I dan Pembimbing II).", indent=False)
add_paragraph(
    "Data UMKM yang tersimpan dalam tabel belum langsung menunjukkan pola kedekatan wilayah dan "
    "tingkat ketelitian lokasinya. Penelitian ini bertujuan merancang Sistem Informasi Geografis "
    "berbasis web untuk pemetaan dan analisis zonasi UMKM di Sulawesi Utara menggunakan K-Means. "
    "Sistem menggunakan React dan Vite untuk antarmuka, Leaflet untuk peta, serta Supabase untuk "
    "basis data dan autentikasi ketika konfigurasi backend tersedia. Analisis menerapkan jarak "
    "Haversine, inisialisasi K-Means++ deterministik, maksimum 100 iterasi, dan ambang konvergensi "
    "0,0001 derajat. Dataset proyek memuat 1.802 UMKM dengan koordinat yang dapat dianalisis. Pada "
    "K=3, algoritma selesai dalam tiga iterasi dan menghasilkan kelompok beranggota 1.116, 46, dan "
    "640 titik dengan WCSS 1.241.741,19 km2. Sistem menyediakan pencarian, filter, informasi kualitas "
    "lokasi, UMKM terdekat, panel admin, verifikasi koordinat, impor-ekspor, audit log, dan penyimpanan "
    "snapshot analisis. Pemisahan koordinat analisis dan tampilan meningkatkan keterlacakan hasil."
)
add_paragraph(
    "Kata kunci: sistem informasi geografis, UMKM, K-Means, Haversine, Leaflet, Supabase.",
    indent=False,
)
next_page()

add_centered("ABSTRACT", 14, True, 20)
add_paragraph(f"{NAME}, {TITLE.title()} (supervised by Advisor I and Advisor II).", indent=False)
add_paragraph(
    "Tabular MSME records do not immediately reveal spatial proximity or location accuracy. This study "
    "designs a web-based Geographic Information System for mapping and zoning MSMEs in North Sulawesi "
    "using K-Means. React and Vite provide the interface, Leaflet provides interactive maps, and "
    "Supabase provides optional database and authentication services. The analysis uses Haversine "
    "distance, deterministic K-Means++ initialization, a maximum of 100 iterations, and a convergence "
    "threshold of 0.0001 degrees. The project dataset contains 1,802 mappable records. With K=3, the "
    "algorithm converged in three iterations and produced groups of 1,116, 46, and 640 records with a "
    "WCSS of 1,241,741.19 km2. The application includes search, category filters, location-quality "
    "labels, nearby-business discovery, administrative data management, coordinate verification, "
    "import-export, audit logs, and reproducible analysis snapshots."
)
add_paragraph(
    "Keywords: geographic information system, MSME, K-Means, Haversine, Leaflet, Supabase.",
    indent=False,
)
next_page()

add_centered("KATA PENGANTAR", 14, True, 20)
add_paragraph(
    "Puji syukur penulis panjatkan kepada Tuhan Yang Maha Kuasa karena atas berkat dan rahmat-Nya "
    "penulis dapat menyelesaikan skripsi ini sebagai salah satu syarat kelulusan Program Sarjana "
    "Terapan pada Program Studi D-IV Teknik Informatika, Jurusan Teknik Elektro, Politeknik Negeri Manado."
)
add_paragraph(
    "Penulis menyampaikan terima kasih kepada pimpinan institusi, dosen pembimbing, seluruh dosen dan "
    "tenaga kependidikan, orang tua serta keluarga, rekan mahasiswa, dan semua pihak yang memberikan "
    "dukungan dalam pelaksanaan penelitian dan penyusunan naskah."
)
add_paragraph(
    "Naskah ini masih memerlukan penyempurnaan identitas akademik, bukti pengujian akhir, validasi "
    "lapangan, dan penyesuaian dengan arahan pembimbing. Semoga hasil penelitian bermanfaat bagi "
    "pengembangan sistem informasi geografis serta pengelolaan data UMKM di Sulawesi Utara."
)
add_centered("Manado, .......................... 2026", after=25)
add_centered("Penulis,", after=55)
add_centered(NAME, 12, True)
next_page()

add_centered("DAFTAR ISI", 14, True, 18)
for entry in (
    "HALAMAN JUDUL ........................................................ i",
    "HALAMAN PENGESAHAN .................................................. ii",
    "PERNYATAAN KEASLIAN ................................................ iii",
    "ABSTRAK ............................................................. iv",
    "ABSTRACT ............................................................. v",
    "KATA PENGANTAR ...................................................... vi",
    "DAFTAR ISI ........................................................ vii",
    "DAFTAR GAMBAR DAN TABEL ............................................. ix",
    "BAB I PENDAHULUAN ..................................................... 1",
    "BAB II TINJAUAN PUSTAKA ............................................... 8",
    "BAB III METODOLOGI PENELITIAN ........................................ 22",
    "BAB IV HASIL DAN PEMBAHASAN .......................................... 35",
    "BAB V PENUTUP ........................................................ 52",
    "DAFTAR PUSTAKA ....................................................... 56",
    "LAMPIRAN ............................................................. 57",
):
    add_paragraph(entry, indent=False)
next_page()

add_centered("DAFTAR GAMBAR", 14, True, 16)
for entry in (
    "Gambar 3.1 Tahapan penelitian",
    "Gambar 3.2 Arsitektur Web GIS Zonasi UMKM",
    "Gambar 3.3 Alur pengolahan data",
    "Gambar 3.4 Alur algoritma K-Means",
    "Gambar 4.1 Alur halaman publik",
    "Gambar 4.2 Alur verifikasi koordinat",
):
    add_paragraph(entry, indent=False)
add_centered("DAFTAR TABEL", 14, True, 16)
for entry in (
    "Tabel 3.1 Kebutuhan fungsional",
    "Tabel 3.2 Kebutuhan nonfungsional",
    "Tabel 4.1 Ringkasan dataset",
    "Tabel 4.2 Matriks peran pengguna",
    "Tabel 4.3 Hasil K-Means",
    "Tabel 4.4 Skenario pengujian",
):
    add_paragraph(entry, indent=False)
next_page()

chapters = {
    "BAB I\nPENDAHULUAN": [
        ("1.1 Latar Belakang", "UMKM berperan dalam kegiatan ekonomi lokal, namun data administrasi belum langsung memperlihatkan pola sebarannya.", "integrasi tabel UMKM dengan peta interaktif", "peta membantu masyarakat dan pengelola membaca kedekatan serta konsentrasi lokasi"),
        ("1.2 Identifikasi Masalah", "Alamat tidak seragam, koordinat perkiraan, dan pembaruan tanpa jejak dapat mengurangi keandalan pembacaan.", "normalisasi, status ketelitian, dan audit perubahan", "pengguna mengetahui batas ketepatan informasi sebelum mengambil tindakan"),
        ("1.3 Rumusan Masalah", "Masalah penelitian mencakup desain Web GIS, penerapan K-Means, kualitas koordinat, dan pengujian fungsi.", "perumusan kebutuhan berdasarkan alur publik dan admin", "pertanyaan penelitian diarahkan pada artefak yang dapat diperiksa"),
        ("1.4 Tujuan Penelitian", "Tujuan penelitian adalah menghasilkan aplikasi yang memadukan informasi usaha, peta, dan zonasi.", "pengembangan fungsi publik, pengelola, dan analisis", "hasil dapat mendukung pencarian UMKM serta eksplorasi pola wilayah"),
        ("1.5 Manfaat Penelitian", "Manfaat akademis berhubungan dengan penerapan algoritma spasial; manfaat praktis berhubungan dengan akses informasi.", "dokumentasi arsitektur dan hasil eksekusi", "penelitian memberi dasar pengembangan sistem lebih lanjut"),
        ("1.6 Batasan Masalah", "K-Means hanya mengolah lintang dan bujur serta tidak menilai omzet, mutu produk, atau kinerja usaha.", "nilai K antara 2 dan 10 pada dataset proyek", "hasil dipakai sebagai dukungan eksplorasi, bukan keputusan tunggal"),
        ("1.7 Sistematika Penulisan", "Naskah dibagi menjadi pendahuluan, teori, metode, hasil, penutup, pustaka, dan lampiran.", "alur masalah-metode-bukti-simpulan", "pembaca dapat membedakan fakta implementasi dan rencana evaluasi"),
    ],
    "BAB II\nTINJAUAN PUSTAKA": [
        ("2.1 Sistem Informasi Geografis", "SIG mengelola data yang memiliki referensi keruangan.", "lapisan titik, basemap, marker, popup, dan panel", "tabel usaha memperoleh konteks geografis"),
        ("2.2 Konsep UMKM", "UMKM menjadi objek data yang memiliki identitas, kategori, alamat, dan lokasi.", "atribut usaha yang tersedia pada dataset", "kualitas atribut menentukan kualitas informasi"),
        ("2.3 Web GIS", "Web GIS menyajikan fungsi geografis melalui browser dan jaringan.", "aplikasi klien React dengan sumber JSON atau Supabase", "pengguna tidak memerlukan perangkat lunak GIS desktop"),
        ("2.4 Data Spasial Titik", "Titik direpresentasikan oleh pasangan lintang dan bujur.", "validasi koordinat analisis dan tampilan", "sumber koordinat tidak tertukar antara analisis dan presentasi"),
        ("2.5 Ketelitian Lokasi", "Status lokasi terdiri dari tepat, perkiraan kecamatan, dan belum terverifikasi.", "label kualitas dan alur verifikasi", "pengguna memahami tingkat keyakinan posisi"),
        ("2.6 Algoritma K-Means", "K-Means mengulang penugasan titik dan pembaruan centroid untuk membentuk K kelompok.", "perhitungan hingga konvergen atau 100 iterasi", "algoritma mudah direproduksi untuk eksplorasi zonasi"),
        ("2.7 Inisialisasi K-Means++", "K-Means++ memilih centroid awal dengan peluang berbobot jarak kuadrat.", "seed deterministik dari K, id, dan koordinat", "masukan sama menghasilkan inisialisasi yang stabil"),
        ("2.8 Jarak Haversine", "Haversine menghitung jarak lengkung bumi dari lintang dan bujur.", "radius bumi 6371,0088 km", "jarak geografis lebih bermakna daripada selisih derajat"),
        ("2.9 WCSS", "WCSS menjumlahkan kuadrat jarak anggota terhadap centroidnya.", "perbandingan WCSS untuk beberapa nilai K", "pemilihan K mempertimbangkan penurunan WCSS dan interpretabilitas"),
        ("2.10 React dan Vite", "React menyusun antarmuka dari komponen dan state, sedangkan Vite menangani pengembangan serta build.", "App, Sidebar, BusinessList, Map, dan halaman admin", "tanggung jawab antarmuka terpisah"),
        ("2.11 Leaflet", "Leaflet menyediakan peta, layer, marker, popup, lingkaran, dan kontrol zoom.", "React-Leaflet dan leaflet.markercluster", "ribuan titik tetap dapat dieksplorasi"),
        ("2.12 Supabase", "Supabase menyediakan basis data, autentikasi, dan Data API ketika backend diaktifkan.", "publishable key di klien, RLS, RPC, dan profil peran", "operasi sensitif tidak bergantung pada tombol antarmuka"),
        ("2.13 Pengujian Perangkat Lunak", "Black-box menguji perilaku fitur dari masukan dan keluaran yang terlihat.", "skenario fungsi, lint, dan build", "bukti pengujian dapat diperiksa kembali"),
        ("2.14 Posisi Penelitian", "Penelitian menempatkan keterlacakan koordinat dan snapshot analisis sebagai bagian dari Web GIS.", "pemetaan publik, verifikasi, K-Means, dan audit", "sistem menggabungkan eksplorasi dengan tata kelola data"),
    ],
    "BAB III\nMETODOLOGI PENELITIAN": [
        ("3.1 Jenis Penelitian", "Penelitian menggunakan pendekatan rancang bangun perangkat lunak.", "identifikasi, desain, implementasi, uji, dan evaluasi", "hasil berupa aplikasi beserta dokumentasi"),
        ("3.2 Objek Penelitian", "Objek penelitian adalah data UMKM dalam cakupan Sulawesi Utara pada dataset proyek.", "pemeriksaan data dan kode sumber", "cakupan mengikuti titik yang tersedia"),
        ("3.3 Sumber Data", "Sumber primer berupa dataset serta artefak kode; sumber sekunder berupa literatur dan dokumentasi.", "umkm.json, komponen React, layanan, dan migrasi", "asal bukti dapat ditelusuri"),
        ("3.4 Tahapan Penelitian", "Tahapan mengalir dari studi literatur sampai dokumentasi hasil.", "pengumpulan, normalisasi, desain, implementasi, uji, analisis", "keputusan teknis mempunyai konteks"),
        ("3.5 Kebutuhan Fungsional", "Fungsi diturunkan dari kebutuhan masyarakat dan pengelola.", "peta, cari, filter, detail, lokasi terdekat, CRUD, verifikasi, impor, ekspor, K-Means, audit", "setiap peran memperoleh fungsi yang sesuai"),
        ("3.6 Kebutuhan Nonfungsional", "Kebutuhan mencakup keamanan, kinerja, responsivitas, dan keterlacakan.", "RLS, validasi, hash dataset, serta desain responsif", "kualitas tidak diukur dari fitur saja"),
        ("3.7 Arsitektur Sistem", "Arsitektur memisahkan presentasi, logika peta, layanan data, dan backend.", "React, Leaflet, JSON statis, dan Supabase", "halaman publik tetap dapat berjalan tanpa backend"),
        ("3.8 Perancangan Data", "Skema menyimpan identitas usaha dan metadata lokasi.", "koordinat analisis, koordinat tampilan, status, area, aktif, dan publikasi", "validasi pasangan koordinat dapat diterapkan"),
        ("3.9 Perancangan K-Means", "Masukan, inisialisasi, penugasan, pembaruan, kondisi henti, dan keluaran ditetapkan.", "K-Means++ deterministik, Haversine, ambang 0,0001", "parameter dapat disimpan untuk reproduksi"),
        ("3.10 Perancangan Antarmuka", "Antarmuka publik berfokus pada penemuan, sedangkan admin berfokus pada pengelolaan.", "sidebar, daftar, peta, panel info, dashboard, dan formulir", "beban kerja pengguna dipisahkan"),
        ("3.11 Perancangan Keamanan", "Keamanan mencakup autentikasi dan otorisasi baris di basis data.", "Supabase Auth, RLS, dan RPC tervalidasi", "publishable key tidak menjadi pengganti kebijakan akses"),
        ("3.12 Rencana Pengujian", "Pengujian mencakup jalur normal, input salah, izin lokasi, dan hak akses.", "kasus black-box dan pemeriksaan build", "hasil akhir dapat dilengkapi bukti seminar"),
        ("3.13 Risiko Penelitian", "Risiko utama berasal dari kualitas alamat, titik perkiraan, dan perubahan dataset.", "status lokasi, verifikasi manual, snapshot, dan hash", "hasil lama dapat dibandingkan dengan masukan aslinya"),
    ],
    "BAB IV\nHASIL DAN PEMBAHASAN": [
        ("4.1 Gambaran Implementasi", "Aplikasi bernama Zonasi UMKM Sulawesi Utara memiliki halaman publik dan panel /admin.", "React, Vite, Leaflet, marker cluster, dan Supabase opsional", "pencarian publik dan tata kelola admin terpisah"),
        ("4.2 Dataset Proyek", "Dataset memuat 1.802 rekaman yang seluruhnya mempunyai koordinat analisis.", "normalisasi JSON sebelum digunakan", "angka menggambarkan snapshot proyek, bukan statistik resmi provinsi"),
        ("4.3 Komposisi Kategori", "Kategori terbesar adalah Roti, Kue & Biskuit 596; Jasa Boga/Katering 569; Makanan Ringan 173.", "agregasi product_label dan product_type", "komposisi membantu pembaca memahami dominasi dataset"),
        ("4.4 Normalisasi Lokasi", "Data lama dipetakan ke field analisis dan tampilan tanpa mengklaim titik GPS yang presisi.", "normalizeBusinessLocation", "kompatibilitas lama tetap dipertahankan"),
        ("4.5 Halaman Publik", "Sidebar, pencarian, kategori, daftar, dan peta membentuk alur penemuan UMKM.", "state React dan sinkronisasi pilihan", "pengguna tidak perlu membaca tabel panjang"),
        ("4.6 Marker Cluster", "Marker dikelompokkan sesuai zoom dan dibuka ketika pengguna mendekat.", "leaflet.markercluster dengan fallback layer", "peta tetap responsif dan informatif"),
        ("4.7 Informasi Popup", "Popup menampilkan nama, merek, kategori, pemilik, alamat, status lokasi, dan zona.", "escapeHtml dan binding popup", "nilai dataset tidak dieksekusi sebagai markup"),
        ("4.8 UMKM Terdekat", "Sistem mengambil lokasi browser dan mengurutkan lima titik terdekat.", "Geolocation API dan Haversine", "hasil perkiraan diberi label eksplisit"),
        ("4.9 Panel Admin", "Admin memuat dashboard, data, verifikasi, K-Means, impor-ekspor, dan audit.", "routing antarmuka berdasarkan sesi dan profil", "pekerjaan operasional memiliki ruang terpisah"),
        ("4.10 Peran Pengguna", "Superadmin/admin dapat menulis, verifikator terbatas pada verifikasi, dan viewer membaca.", "RLS dan pemeriksaan peran", "menyembunyikan tombol bukan satu-satunya perlindungan"),
        ("4.11 Verifikasi Koordinat", "Titik diperiksa melalui pin atau input angka sebelum berstatus tepat.", "RPC verify_umkm_location dan validasi pasangan", "bukti ketelitian tidak dipulihkan melalui CSV"),
        ("4.12 K-Means", "Algoritma menugaskan titik ke centroid terdekat dan menghitung centroid baru.", "performKMeans dengan Haversine", "keluaran mencakup cluster, centroid, WCSS, iterasi, dan radius"),
        ("4.13 Hasil K-Means", "Pada K=3 diperoleh kelompok 1.116, 46, dan 640 dengan tiga iterasi.", "eksekusi modul proyek pada 1.802 titik", "kelompok kecil perlu dipertimbangkan saat memilih K"),
        ("4.14 Perbandingan Nilai K", "WCSS turun dari 3.716.836,85 pada K=2 menjadi 253.724,57 pada K=6.", "perhitungan konsisten pada snapshot yang sama", "penurunan WCSS tidak cukup tanpa interpretasi"),
        ("4.15 Snapshot Analisis", "Masukan diurutkan, disimpan, dan diberi hash SHA-256 bersama parameter.", "kmeans_runs, input_snapshot, dan dataset_hash", "hasil historis dapat direproduksi"),
        ("4.16 Impor dan Audit", "Impor batch atomik mencegah penyimpanan sebagian ketika validasi gagal.", "RPC import_umkm_batch dan audit_logs", "perubahan data mempunyai jejak"),
        ("4.17 Pengujian dan Keterbatasan", "Skenario menguji fungsi publik, admin, validasi, dan akses; kualitas titik tetap menjadi batas utama.", "black-box, lint, build, dan verifikasi lapangan", "hasil sistem perlu dibaca sesuai mutu sumber"),
    ],
    "BAB V\nPENUTUP": [
        ("5.1 Kesimpulan Perancangan", "Web GIS berhasil memadukan dataset UMKM, peta, dan analisis zonasi.", "komponen React dan Leaflet", "sistem menyediakan alur pencarian yang dapat digunakan"),
        ("5.2 Kesimpulan Analisis", "K-Means++ dan Haversine menghasilkan klaster yang dapat diulang pada snapshot sama.", "parameter, WCSS, iterasi, centroid, dan hash", "hasil dapat ditelusuri kembali"),
        ("5.3 Saran Teknis", "Standarisasi alamat dan verifikasi titik perlu diprioritaskan.", "geocoding tervalidasi, tanggal pembaruan, dan uji kinerja", "nilai informasi spasial meningkat"),
        ("5.4 Saran Penelitian", "Evaluasi usability dan validasi pemangku kepentingan perlu dilakukan sebelum pengajuan akhir.", "responden terencana dan bukti pengujian", "simpulan empiris memperoleh dukungan yang memadai"),
    ],
}

for chapter, topics in chapters.items():
    for section_title, evidence, implementation, interpretation in topics:
        add_topic_page(
            chapter,
            section_title,
            academic_paragraphs(section_title, evidence, implementation, interpretation),
        )

add_heading("DAFTAR PUSTAKA")
for reference in (
    "Arthur, D., & Vassilvitskii, S. (2007). k-means++: The Advantages of Careful Seeding. Proceedings of SODA.",
    "Badan Pusat Statistik. (2024). Gambaran Usaha di Indonesia. Jakarta: BPS.",
    "Leaflet. (2026). Leaflet API Reference 1.9.4. https://leafletjs.com/reference.",
    "MacQueen, J. (1967). Some Methods for Classification and Analysis of Multivariate Observations. Proceedings of the Fifth Berkeley Symposium.",
    "OpenStreetMap contributors. (2026). OpenStreetMap. https://www.openstreetmap.org.",
    "React. (2026). React Documentation. https://react.dev/learn.",
    "Sinnott, R. W. (1984). Virtues of the Haversine. Sky and Telescope, 68(2), 159.",
    "Supabase. (2026). Documentation: Auth, Data API, and Row Level Security. https://supabase.com/docs.",
    "Proyek Web GIS Zonasi UMKM Sulawesi Utara. (2026). Kode sumber dan dokumentasi proyek.",
):
    add_paragraph(reference, indent=False)
next_page()

syntax_pages = [
    (
        "LAMPIRAN 1\nPENJELASAN SINTAKS K-MEANS",
        [
            ["const EARTH_RADIUS_KM = 6371.0088;", "Menyimpan radius rata-rata bumi dalam kilometer untuk rumus Haversine."],
            ["Number(point.analysis_lat ?? point.lat)", "Memilih koordinat analisis; field lama dipakai hanya jika field baru null atau undefined."],
            ["Math.PI / 180", "Mengubah derajat menjadi radian sebelum fungsi trigonometri dijalankan."],
            ["Math.sin(dLat / 2) ** 2", "Menghitung komponen perubahan lintang pada rumus Haversine."],
            ["Math.atan2(Math.sqrt(a), Math.sqrt(1 - a))", "Menghasilkan sudut pusat yang stabil secara numerik."],
            ["createSeed(data, k)", "Membentuk seed dari isi dataset dan K agar hasil inisialisasi dapat diulang."],
            ["seededRandom(seed)", "Menghasilkan urutan bilangan semu acak yang konsisten untuk seed sama."],
            ["distances.map(d => d / totalDist)", "Mengubah jarak kuadrat menjadi probabilitas pemilihan centroid K-Means++."],
            ["new Array(k).fill().map(() => [])", "Membuat K array cluster yang terpisah, bukan referensi ke array yang sama."],
            ["iterations < maxIterations", "Mencegah proses iterasi berjalan tanpa batas."],
        ],
    ),
    (
        "LAMPIRAN 2\nPENJELASAN SINTAKS HASIL K-MEANS",
        [
            ["for (const point of data)", "Mengunjungi setiap titik masukan untuk proses penugasan cluster."],
            ["distance(point, centroids[i])", "Mengukur jarak geografis titik terhadap centroid ke-i."],
            ["clusters[closestCentroidIndex].push(point)", "Memasukkan titik ke cluster dengan jarak minimum."],
            ["cluster.reduce(...)", "Menjumlahkan lintang atau bujur seluruh anggota cluster."],
            ["sumLat / cluster.length", "Menghitung lintang centroid baru sebagai rata-rata anggota."],
            ["Math.abs(old - current) > 0.0001", "Menentukan apakah perpindahan centroid masih signifikan."],
            ["Math.pow(distance(...), 2)", "Menambahkan kuadrat jarak ke nilai WCSS."],
            ["calculateClusterRadius", "Mencari jarak terjauh anggota terhadap centroid untuk visualisasi radius."],
            ["return { ...point, cluster: clusterIndex }", "Menghasilkan salinan titik dengan label cluster tanpa mengubah objek sumber."],
            ["generateClusterColors(k)", "Menetapkan warna berbeda agar cluster dapat dibedakan pada peta."],
        ],
    ),
    (
        "LAMPIRAN 3\nPENJELASAN SINTAKS LOKASI",
        [
            ["finiteNumber(value)", "Mengubah nilai koordinat menjadi angka valid atau null."],
            ["Number.isFinite(number)", "Menolak NaN dan Infinity sebagai koordinat."],
            ["analysis_lat ?? lat", "Menggunakan field lama hanya ketika koordinat analisis belum tersedia."],
            ["location_accuracy || APPROXIMATE", "Memberi status aman pada data lama yang berasal dari pencocokan wilayah."],
            ["getAnalysisCoordinates", "Mengambil titik untuk perhitungan jarak dan K-Means."],
            ["getDisplayCoordinates", "Mengambil titik yang dipakai untuk posisi marker di peta."],
            ["isMappableLocation", "Memastikan status bukan belum terverifikasi dan pasangan koordinat tersedia."],
            ["locationAccuracyLabel", "Mengubah nilai status internal menjadi teks yang dipahami pengguna."],
        ],
    ),
    (
        "LAMPIRAN 4\nPENJELASAN SINTAKS PEMUATAN DATA",
        [
            ["import.meta.env.VITE_SUPABASE_URL", "Membaca URL proyek Supabase yang diekspos Vite ke frontend."],
            ["Boolean(url && key)", "Menentukan apakah konfigurasi backend lengkap sebelum membuat permintaan."],
            ["fetch('/data/umkm.json')", "Memuat dataset statis ketika Supabase belum dikonfigurasi."],
            ["if (!response.ok)", "Menghentikan alur ketika server mengembalikan status HTTP gagal."],
            ["Array.isArray(data)", "Memastikan struktur JSON berbentuk daftar UMKM."],
            ["data.map(normalizeBusinessLocation)", "Menormalisasi seluruh objek ke struktur koordinat yang sama."],
            ["await import('./umkmService')", "Memuat modul Supabase hanya ketika dibutuhkan."],
            ["query.range(from, to)", "Mengambil data secara bertahap agar tidak bergantung pada satu respons tak terbatas."],
        ],
    ),
    (
        "LAMPIRAN 5\nPENJELASAN SINTAKS PETA",
        [
            ["<MapContainer center={...} zoom={8}>", "Membuat konteks peta Leaflet dengan pusat dan zoom awal."],
            ["<TileLayer url={...} />", "Menampilkan basemap raster beserta atribusinya."],
            ["L.markerClusterGroup({...})", "Mengelompokkan marker berdekatan agar render lebih ringan."],
            ["marker.bindPopup(content)", "Menghubungkan detail UMKM dengan marker yang dipilih."],
            ["map.flyTo([lat, lng], zoom)", "Memindahkan tampilan peta secara terarah ke lokasi terpilih."],
            ["navigator.geolocation.getCurrentPosition", "Meminta posisi perangkat setelah persetujuan pengguna."],
            ["sort((a, b) => a.distance - b.distance)", "Mengurutkan hasil dari jarak terkecil."],
            ["slice(0, 5)", "Membatasi panel pada lima UMKM terdekat."],
            ["encodeURIComponent(destination)", "Mengamankan tujuan ketika dimasukkan ke URL rute."],
        ],
    ),
    (
        "LAMPIRAN 6\nPENJELASAN SINTAKS SUPABASE",
        [
            ["supabase.from('umkm').select(fields)", "Membaca kolom UMKM melalui Data API sesuai kebijakan akses."],
            [".eq('is_active', true)", "Membatasi hasil pada data yang masih aktif."],
            [".eq('published', true)", "Membatasi halaman publik pada data yang diizinkan tampil."],
            [".update(payload).eq('id', id)", "Memperbarui baris tertentu; RLS tetap menentukan apakah operasi diizinkan."],
            ["supabase.rpc('verify_umkm_location', args)", "Menjalankan prosedur database untuk alur verifikasi yang dibatasi."],
            ["supabase.rpc('import_umkm_batch', args)", "Mengirim satu batch agar validasi dan penulisan dapat berlangsung atomik."],
            ["publishable key", "Kunci yang boleh digunakan frontend; bukan pengganti RLS dan bukan service-role key."],
            ["Row Level Security", "Membatasi baris dan operasi berdasarkan kebijakan database serta identitas pengguna."],
            ["USING dan WITH CHECK", "USING membatasi baris yang dapat dipilih/diubah; WITH CHECK memvalidasi nilai baru pada UPDATE."],
        ],
    ),
    (
        "LAMPIRAN 7\nPENJELASAN SINTAKS SNAPSHOT DAN HASH",
        [
            ["analysisData.filter(...) ", "Memilih data aktif, dipublikasikan, dan mempunyai lokasi analitis."],
            [".map(({id, lat, lng}) => ...)", "Membentuk snapshot minimal yang cukup untuk mengulang analisis."],
            [".sort(compareSnapshotPoints)", "Menetapkan urutan stabil sebelum hash dihitung."],
            ["JSON.stringify(value)", "Mengubah snapshot menjadi representasi byte yang konsisten."],
            ["crypto.subtle.digest('SHA-256', bytes)", "Menghasilkan sidik data untuk mendeteksi perubahan masukan."],
            ["datasetHash.slice(0, 16)", "Menampilkan ringkasan hash tanpa menghilangkan hash penuh yang disimpan."],
            ["saveKMeansRun({...})", "Menyimpan K, jumlah data, iterasi, WCSS, centroid, statistik, snapshot, dan parameter."],
            ["order('created_at', { ascending: false })", "Menampilkan riwayat terbaru terlebih dahulu."],
            ["limit(30)", "Membatasi jumlah riwayat pada tampilan agar respons tetap terkontrol."],
        ],
    ),
]

for index, (title, rows) in enumerate(syntax_pages):
    add_heading(title)
    add_paragraph(
        "Tabel berikut memberi komentar fungsi untuk setiap sintaks penting pada potongan kode inti. "
        "Penjelasan mengikuti perilaku aktual di proyek dan tidak mengubah source code produksi.",
        indent=False,
    )
    add_table(["Sintaks", "Komentar fungsi"], rows, [6.2, 8.0])
    if index < len(syntax_pages) - 1:
        next_page()

settings = document.settings.element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)
document.core_properties.title = "Skripsi Web GIS Zonasi UMKM Sulawesi Utara"
document.core_properties.author = NAME
document.core_properties.subject = "Web GIS dan K-Means untuk pemetaan UMKM"
document.save(OUTPUT)
print(OUTPUT)
